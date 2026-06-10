"""SPC 분석 보고서 자동 생성 (Excel + Word)."""
from __future__ import annotations

import logging
from datetime import datetime
from pathlib import Path
from typing import Optional

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from src.spc.statistics import SpcAnalysisResult

logger = logging.getLogger(__name__)


class SpcReportGenerator:
    """분석 결과를 Excel·Word 보고서로 출력."""

    def __init__(self, output_dir: str | Path):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def generate(
        self,
        result: SpcAnalysisResult,
        *,
        raw_sample: pd.DataFrame,
        chart_path: Optional[Path] = None,
        histogram_path: Optional[Path] = None,
        report_title: str = "SPC 공정능력 분석 보고서",
        metadata: Optional[dict] = None,
    ) -> dict[str, Path]:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base = self.output_dir / f"spc_report_{timestamp}"

        excel_path = self._write_excel(base.with_suffix(".xlsx"), result, raw_sample, metadata)
        word_path = self._write_word(
            base.with_suffix(".docx"),
            result,
            report_title,
            chart_path,
            histogram_path,
            metadata,
        )

        logger.info("보고서 생성 완료: %s, %s", excel_path, word_path)
        return {"excel": excel_path, "word": word_path}

    def _write_excel(
        self,
        path: Path,
        result: SpcAnalysisResult,
        raw_sample: pd.DataFrame,
        metadata: Optional[dict],
    ) -> Path:
        with pd.ExcelWriter(path, engine="openpyxl") as writer:
            # 요약
            summary_rows = []
            summary_rows.append(result.normality.to_dict())
            summary_rows.append(result.control_limits.to_dict())
            if result.capability:
                summary_rows.append(result.capability.to_dict())
            summary_rows.append({
                "관리상태": "관리外" if result.out_of_control_points else "관리內",
                "이탈점": ", ".join(map(str, result.out_of_control_points)) or "없음",
            })
            if metadata:
                summary_rows.append(metadata)

            pd.DataFrame(summary_rows).to_excel(writer, sheet_name="요약", index=False)

            # 정규성
            pd.DataFrame([result.normality.to_dict()]).to_excel(writer, sheet_name="정규성검정", index=False)

            # 관리도 한계
            limits_data = []
            cl = result.control_limits
            if cl.xbar_limits:
                for k, v in cl.xbar_limits.items():
                    limits_data.append({"차트": "X-bar", "항목": k, "값": v})
            if cl.r_limits:
                for k, v in cl.r_limits.items():
                    limits_data.append({"차트": "R", "항목": k, "값": v})
            if cl.i_limits:
                for k, v in cl.i_limits.items():
                    limits_data.append({"차트": "I", "항목": k, "값": v})
            if cl.mr_limits:
                for k, v in cl.mr_limits.items():
                    limits_data.append({"차트": "MR", "항목": k, "값": v})
            pd.DataFrame(limits_data).to_excel(writer, sheet_name="관리한계", index=False)

            # 공정능력
            if result.capability:
                pd.DataFrame([result.capability.to_dict()]).to_excel(writer, sheet_name="공정능력", index=False)

            # subgroup / individual 데이터
            if result.subgroup_stats is not None:
                result.subgroup_stats.to_excel(writer, sheet_name="Subgroup통계", index=False)
            if result.individual_stats is not None:
                result.individual_stats.to_excel(writer, sheet_name="Individual통계", index=False)

            raw_sample.to_excel(writer, sheet_name="채취표본", index=False)

        return path

    def _write_word(
        self,
        path: Path,
        result: SpcAnalysisResult,
        title: str,
        chart_path: Optional[Path],
        histogram_path: Optional[Path],
        metadata: Optional[dict],
    ) -> Path:
        doc = Document()
        section = doc.sections[0]

        # 표지
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(18)

        doc.add_paragraph(f"작성일시: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        if metadata:
            for k, v in metadata.items():
                doc.add_paragraph(f"{k}: {v}")

        doc.add_heading("1. 분석 개요", level=1)
        doc.add_paragraph(
            f"관리도 유형: {result.control_limits.chart_type} | "
            f"표본수: {result.normality.n} | "
            f"관리상태: {'관리外 (이탈점 존재)' if result.out_of_control_points else '관리內'}"
        )
        if result.out_of_control_points:
            doc.add_paragraph(f"관리한계 이탈 subgroup/point: {result.out_of_control_points}")

        doc.add_heading("2. 정규성 검정", level=1)
        norm = result.normality
        doc.add_paragraph(
            f"{norm.test_name} 검정 결과 — 통계량: {norm.statistic:.4f}, "
            f"p-value: {norm.p_value:.6f}, 판정: {'정규분포 (α={})'.format(norm.alpha) if norm.is_normal else '비정규'}"
        )

        doc.add_heading("3. 관리도 한계", level=1)
        cl = result.control_limits
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "항목"
        hdr[1].text = "값"
        hdr[2].text = "비고"

        rows = [
            ("중심선 (CL)", f"{cl.center_line:.6f}", ""),
            ("UCL", f"{cl.ucl:.6f}", ""),
            ("LCL", f"{cl.lcl:.6f}", ""),
            ("추정 σ (within)", f"{cl.sigma_estimate:.6f}", "공정능력 산출용"),
        ]
        for label, val, note in rows:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = val
            row[2].text = note

        if chart_path and chart_path.exists():
            doc.add_heading("4. 관리도", level=1)
            doc.add_picture(str(chart_path), width=Inches(6.0))

        doc.add_heading("5. 공정능력 (Process Capability)", level=1)
        if result.capability:
            cap = result.capability
            cap_table = doc.add_table(rows=1, cols=2)
            cap_table.style = "Table Grid"
            cap_table.rows[0].cells[0].text = "지표"
            cap_table.rows[0].cells[1].text = "값"

            cap_items = [
                ("USL", f"{cap.usl}"),
                ("LSL", f"{cap.lsl}"),
                ("평균 (X̄)", f"{cap.mean:.6f}"),
                ("Cp", f"{cap.cp:.4f}"),
                ("Cpk", f"{cap.cpk:.4f}"),
                ("Pp", f"{cap.pp:.4f}"),
                ("Ppk", f"{cap.ppk:.4f}"),
                ("예상 PPM", f"{cap.ppm_est:.2f}"),
            ]
            for label, val in cap_items:
                row = cap_table.add_row().cells
                row[0].text = label
                row[1].text = val

            doc.add_paragraph(self._capability_interpretation(cap.cpk))

            if histogram_path and histogram_path.exists():
                doc.add_paragraph("")
                doc.add_picture(str(histogram_path), width=Inches(5.5))
        else:
            doc.add_paragraph("USL/LSL 미지정으로 공정능력 지표를 산출하지 않았습니다.")

        doc.add_heading("6. 결론", level=1)
        doc.add_paragraph(self._conclusion(result))

        doc.save(path)
        return path

    @staticmethod
    def _capability_interpretation(cpk: float) -> str:
        if cpk >= 1.67:
            level = "우수 (6σ 수준 근접)"
        elif cpk >= 1.33:
            level = "양호 (4σ 수준)"
        elif cpk >= 1.0:
            level = "최소 기준 충족 (3σ 수준)"
        else:
            level = "개선 필요 (공정 변동 또는 평균 이탈)"
        return f"Cpk={cpk:.4f} → {level}"

    @staticmethod
    def _conclusion(result: SpcAnalysisResult) -> str:
        from src.spc.chart_expert_review import format_conclusions_for_report

        return format_conclusions_for_report(result)
